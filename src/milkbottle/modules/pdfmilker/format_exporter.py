"""Multi-format export system for PDFmilker.

This module provides comprehensive export capabilities for extracted PDF content
to multiple formats including Markdown, HTML, LaTeX, JSON, and Word documents.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from rich.console import Console

logger = logging.getLogger(__name__)
console = Console()


class FormatExporter:
    """Multi-format exporter for PDF content."""

    def __init__(self):
        """Initialize the format exporter."""
        self.supported_formats = ["markdown", "html", "latex", "json", "docx"]
        self.templates_dir = Path(__file__).parent / "templates"
        self.templates_dir.mkdir(exist_ok=True)

    def export_to_format(
        self,
        content: Dict[str, Any],
        format_type: str,
        output_path: Path,
        template_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export content to the specified format.

        Args:
            content: Extracted content dictionary
            format_type: Target format (markdown, html, latex, json, docx)
            output_path: Output file path
            template_name: Optional template name for custom formatting

        Returns:
            Export result dictionary
        """
        try:
            format_type = format_type.lower()

            if format_type not in self.supported_formats:
                raise ValueError(f"Unsupported format: {format_type}")

            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)

            if format_type == "markdown":
                result = self._export_markdown(content, output_path, template_name)
            elif format_type == "html":
                result = self._export_html(content, output_path, template_name)
            elif format_type == "latex":
                result = self._export_latex(content, output_path, template_name)
            elif format_type == "json":
                result = self._export_json(content, output_path)
            elif format_type == "docx":
                result = self._export_docx(content, output_path, template_name)
            else:
                raise ValueError(f"Format {format_type} not implemented")

            logger.info(f"Successfully exported to {format_type}: {output_path}")
            return {
                "success": True,
                "output_path": output_path,
                "format": format_type,
                "file_size": output_path.stat().st_size if output_path.exists() else 0,
                "export_details": result,
            }

        except Exception as e:
            logger.error(f"Export to {format_type} failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": format_type,
                "output_path": output_path,
            }

    def _export_markdown(
        self,
        content: Dict[str, Any],
        output_path: Path,
        template_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export content to Markdown format.

        Args:
            content: Extracted content
            output_path: Output file path
            template_name: Optional template name

        Returns:
            Export details
        """
        try:
            # Use template if specified
            if template_name:
                template = self._load_template("markdown", template_name)
                markdown_content = self._apply_template(template, content)
            else:
                markdown_content = self._generate_markdown(content)

            # Write to file
            output_path.write_text(markdown_content, encoding="utf-8")

            return {
                "content_length": len(markdown_content),
                "sections": self._count_sections(content),
                "template_used": template_name or "default",
            }

        except Exception as e:
            logger.error(f"Markdown export failed: {e}")
            raise

    def _export_html(
        self,
        content: Dict[str, Any],
        output_path: Path,
        template_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export content to HTML format.

        Args:
            content: Extracted content
            output_path: Output file path
            template_name: Optional template name

        Returns:
            Export details
        """
        try:
            # Use template if specified
            if template_name:
                template = self._load_template("html", template_name)
                html_content = self._apply_template(template, content)
            else:
                html_content = self._generate_html(content)

            # Write to file
            output_path.write_text(html_content, encoding="utf-8")

            return {
                "content_length": len(html_content),
                "sections": self._count_sections(content),
                "template_used": template_name or "default",
                "has_math": bool(content.get("math_formulas")),
                "has_tables": bool(content.get("tables")),
            }

        except Exception as e:
            logger.error(f"HTML export failed: {e}")
            raise

    def _export_latex(
        self,
        content: Dict[str, Any],
        output_path: Path,
        template_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export content to LaTeX format.

        Args:
            content: Extracted content
            output_path: Output file path
            template_name: Optional template name

        Returns:
            Export details
        """
        try:
            # Use template if specified
            if template_name:
                template = self._load_template("latex", template_name)
                latex_content = self._apply_template(template, content)
            else:
                latex_content = self._generate_latex(content)

            # Write to file
            output_path.write_text(latex_content, encoding="utf-8")

            return {
                "content_length": len(latex_content),
                "sections": self._count_sections(content),
                "template_used": template_name or "default",
                "math_formulas": len(content.get("math_formulas", [])),
                "tables": len(content.get("tables", [])),
            }

        except Exception as e:
            logger.error(f"LaTeX export failed: {e}")
            raise

    def _export_json(
        self, content: Dict[str, Any], output_path: Path
    ) -> Dict[str, Any]:
        """Export content to JSON format.

        Args:
            content: Extracted content
            output_path: Output file path

        Returns:
            Export details
        """
        try:
            # Add metadata
            export_data = {
                "metadata": {
                    "export_format": "json",
                    "export_timestamp": self._get_timestamp(),
                    "content_sections": list(content.keys()),
                },
                "content": content,
            }

            # Write to file
            with output_path.open("w", encoding="utf-8") as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

            return {
                "content_length": len(json.dumps(export_data)),
                "sections": len(content),
                "metadata_included": True,
            }

        except Exception as e:
            logger.error(f"JSON export failed: {e}")
            raise

    def _export_docx(
        self,
        content: Dict[str, Any],
        output_path: Path,
        template_name: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export content to Word document format.

        Args:
            content: Extracted content
            output_path: Output file path
            template_name: Optional template name

        Returns:
            Export details
        """
        try:
            # Import here to avoid dependency issues
            try:
                import importlib.util

                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                if not importlib.util.find_spec("docx.shared"):
                    raise ImportError("docx.shared module not available")
            except ImportError as e:
                raise ImportError(
                    "python-docx is required for Word export. Install with: pip install python-docx"
                ) from e

            # Create document
            doc = Document()

            # Add title
            if content.get("title"):
                title = doc.add_heading(content["title"], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add abstract
            if content.get("abstract"):
                doc.add_heading("Abstract", level=1)
                doc.add_paragraph(content["abstract"])

            # Add body text
            if content.get("body_text"):
                doc.add_heading("Content", level=1)
                doc.add_paragraph(content["body_text"])

            # Add mathematical formulas
            if content.get("math_formulas"):
                doc.add_heading("Mathematical Formulas", level=1)
                for formula in content["math_formulas"]:
                    p = doc.add_paragraph()
                    p.add_run(f"Formula {formula.get('number', 'N/A')}: ")
                    p.add_run(formula.get("latex", formula.get("text", "")))

            # Add tables
            if content.get("tables"):
                doc.add_heading("Tables", level=1)
                for i, table_data in enumerate(content["tables"]):
                    doc.add_paragraph(f"Table {i+1}:")
                    if isinstance(table_data, dict) and "markdown" in table_data:
                        doc.add_paragraph(table_data["markdown"])
                    else:
                        doc.add_paragraph(str(table_data))

            # Add references
            if content.get("references"):
                doc.add_heading("References", level=1)
                for ref in content["references"]:
                    doc.add_paragraph(ref, style="List Bullet")

            # Save document
            doc.save(output_path)

            return {
                "content_length": output_path.stat().st_size,
                "sections": self._count_sections(content),
                "template_used": template_name or "default",
                "has_math": bool(content.get("math_formulas")),
                "has_tables": bool(content.get("tables")),
            }

        except Exception as e:
            logger.error(f"Word export failed: {e}")
            raise

    def _generate_markdown(self, content: Dict[str, Any]) -> str:
        """Generate Markdown content from extracted data.

        Args:
            content: Extracted content

        Returns:
            Markdown string
        """
        lines = []

        # Title
        if content.get("title"):
            lines.extend((f"# {content['title']}", ""))
        # Abstract
        if content.get("abstract"):
            lines.extend(("## Abstract", "", content["abstract"], ""))
        # Body text
        if content.get("body_text"):
            lines.extend(("## Content", "", content["body_text"], ""))
        # Mathematical formulas
        if content.get("math_formulas"):
            lines.extend(("## Mathematical Formulas", ""))
            for formula in content["math_formulas"]:
                lines.extend(
                    (
                        f"**Formula {formula.get('number', 'N/A')}:**",
                        "```latex",
                        formula.get("latex", formula.get("text", "")),
                        "```",
                        "",
                    )
                )
        # Tables
        if content.get("tables"):
            lines.extend(("## Tables", ""))
            for i, table_data in enumerate(content["tables"]):
                lines.extend((f"**Table {i + 1}:**", ""))
                if isinstance(table_data, dict) and "markdown" in table_data:
                    lines.append(table_data["markdown"])
                else:
                    lines.append(str(table_data))
                lines.append("")

        # References
        if content.get("references"):
            lines.extend(("## References", ""))
            lines.extend(f"- {ref}" for ref in content["references"])
            lines.append("")

        return "\n".join(lines)

    def _generate_html(self, content: Dict[str, Any]) -> str:
        """Generate HTML content from extracted data.

        Args:
            content: Extracted content

        Returns:
            HTML string
        """
        html_parts = [
            "<!DOCTYPE html>",
            "<html lang='en'>",
            "<head>",
            "    <meta charset='UTF-8'>",
            "    <meta name='viewport' content='width=device-width, initial-scale=1.0'>",
            "    <title>Extracted PDF Content</title>",
            "    <style>",
            "        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }",
            "        h1 { color: #2c3e50; border-bottom: 2px solid #3498db; }",
            "        h2 { color: #34495e; margin-top: 30px; }",
            "        .math { background-color: #f8f9fa; padding: 10px; border-left: 4px solid #3498db; }",
            "        .table { margin: 20px 0; }",
            "        .reference { margin: 5px 0; }",
            "    </style>",
            "</head>",
            "<body>",
        ]

        # Title
        if content.get("title"):
            html_parts.append(f"<h1>{content['title']}</h1>")

        # Abstract
        if content.get("abstract"):
            html_parts.extend(("<h2>Abstract</h2>", f"<p>{content['abstract']}</p>"))
        # Body text
        if content.get("body_text"):
            html_parts.extend(("<h2>Content</h2>", f"<p>{content['body_text']}</p>"))
        # Mathematical formulas
        if content.get("math_formulas"):
            html_parts.append("<h2>Mathematical Formulas</h2>")
            for formula in content["math_formulas"]:
                html_parts.extend(
                    (
                        "<div class='math'>",
                        f"<strong>Formula {formula.get('number', 'N/A')}:</strong>",
                        f"<pre>{formula.get('latex', formula.get('text', ''))}</pre>",
                        "</div>",
                    )
                )
        # Tables
        if content.get("tables"):
            html_parts.append("<h2>Tables</h2>")
            for i, table_data in enumerate(content["tables"]):
                html_parts.extend(("<div class='table'>", f"<h3>Table {i + 1}</h3>"))
                if isinstance(table_data, dict) and "markdown" in table_data:
                    html_parts.append(f"<pre>{table_data['markdown']}</pre>")
                else:
                    html_parts.append(f"<pre>{str(table_data)}</pre>")
                html_parts.append("</div>")

        # References
        if content.get("references"):
            html_parts.extend(("<h2>References</h2>", "<ul>"))
            html_parts.extend(
                f"<li class='reference'>{ref}</li>" for ref in content["references"]
            )
            html_parts.append("</ul>")

        html_parts.extend(["</body>", "</html>"])

        return "\n".join(html_parts)

    def _generate_latex(self, content: Dict[str, Any]) -> str:
        """Generate LaTeX content from extracted data.

        Args:
            content: Extracted content

        Returns:
            LaTeX string
        """
        latex_parts = [
            "\\documentclass{article}",
            "\\usepackage[utf8]{inputenc}",
            "\\usepackage{amsmath}",
            "\\usepackage{amsfonts}",
            "\\usepackage{amssymb}",
            "\\usepackage{graphicx}",
            "\\usepackage{geometry}",
            "\\geometry{margin=1in}",
            "",
            "\\begin{document}",
            "",
        ]

        # Title
        if content.get("title"):
            latex_parts.extend((f"\\title{{{content['title']}}}", "\\maketitle", ""))
        # Abstract
        if content.get("abstract"):
            latex_parts.extend(
                ("\\begin{abstract}", content["abstract"], "\\end{abstract}", "")
            )
        # Body text
        if content.get("body_text"):
            latex_parts.extend(("\\section{Content}", content["body_text"], ""))
        # Mathematical formulas
        if content.get("math_formulas"):
            latex_parts.append("\\section{Mathematical Formulas}")
            for formula in content["math_formulas"]:
                latex_parts.extend(
                    (
                        f"\\textbf{{Formula {formula.get('number', 'N/A')}:}}",
                        "\\begin{align*}",
                        formula.get("latex", formula.get("text", "")),
                        "\\end{align*}",
                        "",
                    )
                )
        # Tables
        if content.get("tables"):
            latex_parts.append("\\section{Tables}")
            for i, table_data in enumerate(content["tables"]):
                latex_parts.append(f"\\subsection{{Table {i+1}}}")
                if isinstance(table_data, dict) and "latex" in table_data:
                    latex_parts.append(table_data["latex"])
                else:
                    latex_parts.extend(
                        ("\\begin{verbatim}", str(table_data), "\\end{verbatim}")
                    )
                latex_parts.append("")

        # References
        if content.get("references"):
            latex_parts.extend(("\\section{References}", "\\begin{enumerate}"))
            latex_parts.extend(f"\\item {ref}" for ref in content["references"])
            latex_parts.extend(("\\end{enumerate}", ""))
        latex_parts.extend(["\\end{document}"])

        return "\n".join(latex_parts)

    def _load_template(self, format_type: str, template_name: str) -> str:
        """Load a template for the specified format.

        Args:
            format_type: Format type (markdown, html, latex)
            template_name: Template name

        Returns:
            Template content
        """
        template_path = self.templates_dir / f"{template_name}.{format_type}"

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        return template_path.read_text(encoding="utf-8")

    def _apply_template(self, template: str, content: Dict[str, Any]) -> str:
        """Apply template with content substitution.

        Args:
            template: Template string
            content: Content to substitute

        Returns:
            Processed template
        """
        # Simple template substitution - could be enhanced with a proper templating engine
        result = template

        # Replace common placeholders
        replacements = {
            "{{title}}": content.get("title", ""),
            "{{abstract}}": content.get("abstract", ""),
            "{{body_text}}": content.get("body_text", ""),
            "{{math_formulas}}": self._format_math_formulas(
                content.get("math_formulas", [])
            ),
            "{{tables}}": self._format_tables(content.get("tables", [])),
            "{{references}}": self._format_references(content.get("references", [])),
        }

        for placeholder, value in replacements.items():
            result = result.replace(placeholder, str(value))

        return result

    def _format_math_formulas(self, formulas: List[Dict[str, Any]]) -> str:
        """Format mathematical formulas for template substitution."""
        if not formulas:
            return ""

        formatted = [
            f"Formula {formula.get('number', 'N/A')}: {formula.get('latex', formula.get('text', ''))}"
            for formula in formulas
        ]
        return "\n".join(formatted)

    def _format_tables(self, tables: List[Any]) -> str:
        """Format tables for template substitution."""
        if not tables:
            return ""

        formatted = []
        for i, table in enumerate(tables):
            if isinstance(table, dict) and "markdown" in table:
                formatted.append(f"Table {i+1}:\n{table['markdown']}")
            else:
                formatted.append(f"Table {i+1}:\n{str(table)}")

        return "\n\n".join(formatted)

    def _format_references(self, references: List[str]) -> str:
        """Format references for template substitution."""
        return "\n".join([f"- {ref}" for ref in references]) if references else ""

    def _count_sections(self, content: Dict[str, Any]) -> int:
        """Count the number of content sections."""
        return len(
            [
                k
                for k in content
                if k not in ["metadata", "pdf_path", "extraction_method"]
            ]
        )

    def _get_timestamp(self) -> str:
        """Get current timestamp string."""
        from datetime import datetime

        return datetime.now().isoformat()

    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats.

        Returns:
            List of supported format names
        """
        return self.supported_formats.copy()

    def validate_format(self, format_type: str) -> bool:
        """Validate if a format is supported.

        Args:
            format_type: Format to validate

        Returns:
            True if format is supported
        """
        return format_type.lower() in self.supported_formats
